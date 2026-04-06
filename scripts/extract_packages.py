"""
Extract package data from .docx files in /packages/ directory.
Outputs JSON files to /src/data/packages/ and images to /public/package-images/

Usage: python scripts/extract_packages.py
"""

import os
import re
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx Pillow")
    from docx import Document

# Paths
PROJECT_ROOT = Path(__file__).resolve().parent.parent
PACKAGES_DIR = PROJECT_ROOT / "packages"
OUTPUT_JSON_DIR = PROJECT_ROOT / "src" / "data" / "packages"
OUTPUT_IMG_DIR = PROJECT_ROOT / "public" / "package-images"

# Flight cost estimate per person + $20 markup for 2 adults
FLIGHT_COST_PER_PERSON = 0  # Set to 0, will be added in pricing display
MARKUP_PER_2_ADULTS = 20


def slugify(text: str) -> str:
    """Convert text to URL-friendly slug."""
    text = text.lower().strip()
    text = re.sub(r'[^a-z0-9\s-]', '', text)
    text = re.sub(r'[\s-]+', '-', text)
    return text.strip('-')


def generate_short_name(filename: str) -> str:
    """Generate a short display name from filename."""
    name = Path(filename).stem
    # Remove leading duration like "8D7N - "
    name = re.sub(r'^\d+D\d+N\s*-?\s*', '', name)
    # Remove trailing (1), (2) etc
    name = re.sub(r'\s*\(\d+\)\s*$', '', name)
    # Clean up extra spaces/dashes
    name = re.sub(r'\s*-\s*', ' - ', name).strip(' -')
    return name


def extract_days_nights(paragraphs: list) -> tuple:
    """Extract days and nights from summary itinerary."""
    days = 0
    for p in paragraphs:
        text = p.text.strip()
        match = re.match(r'Day\s+(\d+)', text)
        if match:
            day_num = int(match.group(1))
            if day_num > days:
                days = day_num
    nights = max(0, days - 1)
    return days, nights


def extract_places(summary_items: list) -> list:
    """Extract unique place names from summary itinerary."""
    places = []
    known = {
        'Ha Noi', 'Hanoi', 'Da Nang', 'Ho Chi Minh', 'Phu Quoc',
        'Hoi An', 'Ha Long', 'Ninh Binh', 'Ba Na Hills', 'Sapa',
        'Mekong Delta', 'Cu Chi', 'Hue', 'Marble Mountain'
    }
    for item in summary_items:
        for place in known:
            if place.lower() in item.lower() and place not in places:
                places.append(place)
    return places


def extract_images(doc, package_slug: str) -> list:
    """Extract all images from docx and save to public directory."""
    img_dir = OUTPUT_IMG_DIR / package_slug
    img_dir.mkdir(parents=True, exist_ok=True)

    image_paths = []
    img_index = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_index += 1
            image_data = rel.target_part.blob
            # Determine extension
            content_type = rel.target_part.content_type
            ext = '.jpg'
            if 'png' in content_type:
                ext = '.png'
            elif 'gif' in content_type:
                ext = '.gif'
            elif 'webp' in content_type:
                ext = '.webp'

            filename = f"img-{img_index}{ext}"
            filepath = img_dir / filename
            with open(filepath, 'wb') as f:
                f.write(image_data)

            web_path = f"/package-images/{package_slug}/{filename}"
            image_paths.append(web_path)
            print(f"    Saved: {web_path}")

    return image_paths


def extract_pricing(tables: list) -> list:
    """Extract pricing from tables."""
    pricing = []

    for table in tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        first_cell = rows[0].cells[0].text.strip().lower()

        # Skip non-pricing tables (like guest info)
        if 'agent' in first_cell or 'guest' in first_cell or 'inquiry' in first_cell:
            continue

        # This is a pricing table
        tier_name = rows[0].cells[0].text.strip().replace('\n', ' ')

        # Extract adult price
        adult_price = ""
        kid_prices = []
        hotels = []

        for row in rows:
            cell_text = row.cells[0].text.strip()
            if 'adult' in cell_text.lower():
                # Extract USD amount
                match = re.search(r'(\d+)\s*USD', cell_text)
                if match:
                    base_price = int(match.group(1))
                    # Add $20 markup per 2 adults = $10 per person
                    final_price = base_price + 10
                    adult_price = f"{final_price} USD/person"
            elif 'kid' in cell_text.lower() or 'child' in cell_text.lower():
                kid_prices.append(cell_text.split('(')[0].strip())

            # Check for hotel info
            if len(row.cells) >= 3:
                city = row.cells[0].text.strip()
                hotel = row.cells[1].text.strip()
                room = row.cells[2].text.strip()
                if city.lower() not in ['city', '', 'adult', 'kid', 'child'] and \
                   'star' not in city.lower() and 'usd' not in city.lower():
                    hotels.append({
                        "city": city,
                        "hotel": hotel,
                        "room": room
                    })

        pricing.append({
            "tier": tier_name,
            "adultPrice": adult_price,
            "kidPrices": kid_prices,
            "hotels": hotels,
        })

    return pricing


def parse_docx(filepath: str) -> dict:
    """Parse a single docx file and extract all package data."""
    filename = os.path.basename(filepath)
    print(f"\nProcessing: {filename}")

    doc = Document(filepath)
    paragraphs = doc.paragraphs

    short_name = generate_short_name(filename)
    slug = slugify(short_name)

    # Extract images
    images = extract_images(doc, slug)

    # Parse sections
    summary_itinerary = []
    includes = []
    excludes = []
    notes = []
    validity = ""
    deposit_terms = []
    cancellation = []
    other_terms = []
    detailed_itinerary = []
    payment_methods = []

    current_section = None
    current_day = None
    current_day_details = []

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue

        upper = text.upper()

        # Detect section headers
        if upper.startswith('SUMMARY ITINERARY'):
            current_section = 'summary'
            continue
        elif upper == 'INCLUDE:' or upper == 'INCLUSIONS:' or upper == 'INCLUSION:':
            current_section = 'include'
            continue
        elif upper == 'EXCLUDE:' or upper == 'EXCLUSIONS:' or upper == 'EXCLUSION:':
            current_section = 'exclude'
            continue
        elif upper.startswith('NOTE:') or upper == 'NOTES:':
            current_section = 'note'
            if upper != 'NOTE:' and upper != 'NOTES:':
                notes.append(text.replace('NOTE:', '').strip())
            continue
        elif upper.startswith('VALIDITY'):
            current_section = 'validity'
            validity = text
            continue
        elif upper.startswith('DEPOSIT') or upper.startswith('PAYMENT TERM'):
            current_section = 'deposit'
            continue
        elif upper.startswith('CANCELLATION'):
            current_section = 'cancellation'
            continue
        elif upper.startswith('OTHER TERM'):
            current_section = 'other_terms'
            continue
        elif upper.startswith('PAYMENT METHOD'):
            current_section = 'payment'
            continue
        elif upper == 'ITINERARY' or upper == 'DETAILED ITINERARY':
            current_section = 'itinerary'
            continue

        # Check if this starts a new day in itinerary
        day_match = re.match(r'Day\s+(\d+)\s*[:\-]?\s*(.*)', text, re.IGNORECASE)

        if current_section == 'itinerary' and day_match:
            # Save previous day
            if current_day and current_day_details:
                detailed_itinerary.append({
                    "day": current_day,
                    "title": current_day.split(':', 1)[-1].strip() if ':' in current_day else '',
                    "details": current_day_details
                })
            current_day = text
            current_day_details = []
            continue

        # Append to current section
        if current_section == 'summary':
            if re.match(r'Day\s+\d+', text):
                summary_itinerary.append(text)
        elif current_section == 'include':
            if not text.upper().startswith('EXCLUDE'):
                includes.append(text)
        elif current_section == 'exclude':
            if not text.upper().startswith('NOTE'):
                excludes.append(text)
        elif current_section == 'note':
            notes.append(text)
        elif current_section == 'validity':
            validity += ' ' + text
        elif current_section == 'deposit':
            deposit_terms.append(text)
        elif current_section == 'cancellation':
            cancellation.append(text)
        elif current_section == 'other_terms':
            other_terms.append(text)
        elif current_section == 'payment':
            payment_methods.append(text)
        elif current_section == 'itinerary' and current_day:
            if text and not text.upper().startswith('INCLUDE') and not text.upper().startswith('EXCLUDE'):
                current_day_details.append(text)

    # Save last day
    if current_day and current_day_details:
        detailed_itinerary.append({
            "day": current_day,
            "title": current_day.split(':', 1)[-1].strip() if ':' in current_day else '',
            "details": current_day_details
        })

    # Extract days/nights and places
    days, nights = extract_days_nights(paragraphs)
    places = extract_places(summary_itinerary)

    # Extract pricing from tables
    pricing = extract_pricing(doc.tables)

    # Build package data
    package_data = {
        "id": slug,
        "packageName": short_name,
        "slug": slug,
        "days": days,
        "nights": nights,
        "places": places,
        "images": images,
        "heroImage": images[0] if images else f"/package-images/{slug}/img-1.jpg",
        "summaryItinerary": summary_itinerary,
        "detailedItinerary": detailed_itinerary,
        "includes": includes,
        "excludes": excludes,
        "notes": notes,
        "validity": validity.strip(),
        "pricing": pricing,
        "depositTerms": deposit_terms,
        "cancellation": cancellation,
        "otherTerms": other_terms,
        "paymentMethods": payment_methods,
        "travelWindow": validity.strip(),
        "option": f"{nights}N{days}D - {' | '.join(places[:3])}" if places else f"{nights}N{days}D",
    }

    return package_data


def main():
    """Main extraction function."""
    print("=" * 60)
    print("PACKAGE DOCX EXTRACTOR")
    print("=" * 60)

    # Ensure output directories exist
    OUTPUT_JSON_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_IMG_DIR.mkdir(parents=True, exist_ok=True)

    if not PACKAGES_DIR.exists():
        print(f"ERROR: Packages directory not found: {PACKAGES_DIR}")
        sys.exit(1)

    docx_files = list(PACKAGES_DIR.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {PACKAGES_DIR}")
        sys.exit(1)

    print(f"Found {len(docx_files)} package documents\n")

    all_packages = []

    for filepath in sorted(docx_files):
        try:
            package_data = parse_docx(str(filepath))
            all_packages.append(package_data)

            # Save individual JSON
            json_path = OUTPUT_JSON_DIR / f"{package_data['slug']}.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(package_data, f, indent=2, ensure_ascii=False)
            print(f"  -> Saved: {json_path.name}")

        except Exception as e:
            print(f"  ERROR processing {filepath.name}: {e}")
            import traceback
            traceback.print_exc()

    # Save combined packages list
    combined_path = OUTPUT_JSON_DIR / "_all_packages.json"
    with open(combined_path, 'w', encoding='utf-8') as f:
        json.dump(all_packages, f, indent=2, ensure_ascii=False)

    # Generate packagesList.ts
    generate_packages_list_ts(all_packages)

    print(f"\n{'=' * 60}")
    print(f"EXTRACTION COMPLETE")
    print(f"  Packages: {len(all_packages)}")
    print(f"  JSON files: {OUTPUT_JSON_DIR}")
    print(f"  Images: {OUTPUT_IMG_DIR}")
    print(f"{'=' * 60}")

    # Print summary
    print("\nPackage Summary:")
    print("-" * 60)
    for pkg in all_packages:
        print(f"  {pkg['nights']}N{pkg['days']}D | {pkg['packageName']}")
        print(f"         Places: {', '.join(pkg['places'])}")
        print(f"         Images: {len(pkg['images'])}")
        if pkg['pricing']:
            for p in pkg['pricing']:
                print(f"         Pricing: {p['tier']} - {p['adultPrice']}")
        print()


def generate_packages_list_ts(packages: list):
    """Generate the packagesList.ts file for the homepage."""
    ts_path = PROJECT_ROOT / "src" / "data" / "packagesList.ts"

    items = []
    for pkg in packages:
        hero = pkg['heroImage']
        title = pkg['packageName']
        slug = f"/packages/{pkg['slug']}"
        pricing = ""
        if pkg['pricing']:
            price_text = pkg['pricing'][0].get('adultPrice', '')
            match = re.search(r'(\d+)\s*USD', price_text)
            if match:
                pricing = f"From {match.group(1)} USD"

        days_text = f"{pkg['days']} Days"
        nights_text = f"{pkg['nights']}N/{pkg['days']}D"

        items.append({
            "title": title,
            "slug": slug,
            "image": hero,
            "pricing": pricing or "Custom Quote",
            "days": days_text,
            "nights": nights_text,
        })

    # Write TS file
    lines = [
        "// Auto-generated from package documents",
        "// Run: python scripts/extract_packages.py",
        "",
        "export const trendingPackages = [",
    ]

    for item in items:
        lines.append("  {")
        lines.append(f"    title: '{item['title']}',")
        lines.append(f"    slug: '{item['slug']}',")
        lines.append(f"    image: '{item['image']}',")
        lines.append(f"    pricing: '{item['pricing']}',")
        lines.append("    info: [")
        lines.append(f"      {{ icon: '/icons/calendar.svg', text: '{item['days']}' }},")
        lines.append(f"      {{ icon: '/icons/hotel.svg', text: '{item['nights']}' }},")
        lines.append("    ],")
        lines.append("  },")

    lines.append("];")
    lines.append("")
    lines.append("export const explorePackages = trendingPackages;")
    lines.append("")

    with open(ts_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))

    print(f"\n  -> Generated: {ts_path.name}")


if __name__ == "__main__":
    main()
